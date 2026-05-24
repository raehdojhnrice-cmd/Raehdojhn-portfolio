from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("career-materials")
RESUME_PATH = OUT_DIR / "raehdojhn_resume_template.docx"
COVER_PATH = OUT_DIR / "raehdojhn_cover_letter_template.docx"

INK = RGBColor(10, 10, 10)
MUTED = RGBColor(72, 68, 64)
PAPER = "F2F0EF"
SOFT = "E8E4DC"
RED = RGBColor(255, 51, 102)
RED_HEX = "FF3366"
DARK = "0A0A0A"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8D4CC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color="FF3366", size="8", space="7"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_document_base(doc, compact=False):
    section = doc.sections[0]
    margin = 0.62 if compact else 0.85
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.6 if compact else 10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4 if compact else 7)
    normal.paragraph_format.line_spacing = 1.08 if compact else 1.18

    for style_name, size, color, before, after in [
        ("Title", 27 if compact else 26, INK, 0, 3),
        ("Heading 1", 11.5 if compact else 13, RED, 8, 4),
        ("Heading 2", 10.2 if compact else 11.2, INK, 5, 2),
        ("Heading 3", 9.2 if compact else 10.2, MUTED, 4, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Heading 3"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(8.8 if compact else 10)
        style.paragraph_format.left_indent = Inches(0.24)
        style.paragraph_format.first_line_indent = Inches(-0.12)
        style.paragraph_format.space_after = Pt(2 if compact else 3)
        style.paragraph_format.line_spacing = 1.06 if compact else 1.12

    return section


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(2)
    run = footer.add_run("raehdojhn@gmail.com | Los Angeles / Remote | Portfolio: [your-portfolio-url]")
    set_run_font(run, size=7.8, color=MUTED)


def add_nameplate(doc, subtitle, compact=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("RAEHDOJHN")
    set_run_font(run, size=30 if compact else 28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5 if compact else 7)
    run = p.add_run(subtitle.upper())
    set_run_font(run, size=8.2 if compact else 9, color=RED, bold=True)
    paragraph_border_bottom(p, color=RED_HEX, size="8", space="8")


def add_contact_strip(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    paragraph_shading(p, PAPER)
    pieces = [
        "Los Angeles / Remote",
        "raehdojhn@gmail.com",
        "Portfolio: [your-portfolio-url]",
        "V4: [your-portfolio-url]/v4.html",
    ]
    run = p.add_run("   ".join(pieces))
    set_run_font(run, size=8.2, color=INK, bold=True)


def add_section_title(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    set_run_font(run, size=10.5, color=RED, bold=True)


def add_role(doc, title, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    set_run_font(run, size=9.8, color=INK, bold=True)
    run = p.add_run(f" | {meta}")
    set_run_font(run, size=8.7, color=MUTED, italic=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.7)
    run = p.add_run(text)
    set_run_font(run, size=8.65, color=INK)


def add_small_caps_line(doc, text, fill=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if fill:
        paragraph_shading(p, fill)
    run = p.add_run(text.upper())
    set_run_font(run, size=7.8, color=MUTED, bold=True)


def build_resume():
    doc = Document()
    section = set_document_base(doc, compact=True)
    add_footer(section)

    doc.core_properties.title = "Raehdojhn Resume Template"
    doc.core_properties.subject = "Portfolio-compatible creative technologist resume"
    doc.core_properties.author = "Raehdojhn"

    add_nameplate(doc, "Creative Technologist / Creative Direction / Brand Systems / Technical Design", compact=True)
    add_contact_strip(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        "Classically trained artist and creative technologist from Charlotte, North Carolina, now based in Los Angeles. "
        "Builds at the intersection of design, code, and culture - translating creative vision into scalable systems across "
        "brand worlds, interactive websites, AI-assisted production, audio interfaces, and editorial archives."
    )
    set_run_font(run, size=9.2, color=INK)

    add_small_caps_line(
        doc,
        "Creative Direction | Brand Systems | Technical Design | AI Workflows | Audio Production | Visual Design",
        fill=PAPER,
    )

    add_section_title(doc, "Selected Project Work")
    project_rows = [
        (
            "V4 Archive Site - Chaos Lux Desktop",
            "Design + Code + Music | 2026",
            "Recovered archive shell with desktop windows, project ledger, Prompt Flux, media log, system portals, and an iPod-style music player.",
        ),
        (
            "Portfolio Site - Raehdojhn System",
            "Design + Code | 2026",
            "Editorial portfolio system with alphabetized project index, sandwich menu, light/dark themes, cursor interaction, embedded previews, and deployable static routes.",
        ),
        (
            "Future Archives - Ranaan x CCC",
            "Brand System + Interactive Editorial | 2026",
            "Interactive flip magazine with responsive spreads, page turns, touch gestures, campaign image systems, and archive-world storytelling.",
        ),
        (
            "Nodear* - Brand Vault & EPK Portal",
            "Audio + Identity + AI | 2026",
            "Alter-ego brand portal tying visual identity, EPK logic, listening system, and media archive into a deployable experience.",
        ),
        (
            "AUREA / Jiayou / Seoulstice",
            "Brand Worlds + Web Production | 2026",
            "Retail, fragrance, coffee, and skincare concepts spanning editorial identity, product grids, campaign pages, and interactive browsing patterns.",
        ),
    ]
    for title, meta, desc in project_rows:
        add_role(doc, title, meta)
        add_bullet(doc, desc)

    add_section_title(doc, "Experience")
    add_role(doc, "Independent Creative Technologist / Art Director", "Los Angeles / Remote | [2024-Present]")
    for bullet in [
        "Lead end-to-end creative systems from concept through production: art direction, identity, UI/UX, front-end builds, and AI-assisted asset workflows.",
        "Translate references, moodboards, and brand language into deployable websites, interactive project pages, and reusable visual systems.",
        "Prototype and ship static web experiences using HTML, CSS, JavaScript, React-style component thinking, motion logic, and responsive design.",
    ]:
        add_bullet(doc, bullet)
    add_role(doc, "[Company / Studio / Client Role]", "[Location] | [Dates]")
    for bullet in [
        "[Replace with measurable result: launched X campaign/site/system for Y audience or client.]",
        "[Replace with collaboration scope: worked with founders, artists, creative teams, engineers, or clients to define direction.]",
    ]:
        add_bullet(doc, bullet)

    add_section_title(doc, "Skills & Tools")
    skills_table = doc.add_table(rows=4, cols=2)
    set_table_width(skills_table, [1.25, 5.0])
    skills = [
        ("Design", "Visual Identity, Art Direction, Editorial Design, Motion Design, Campaign Systems"),
        ("Development", "React / JavaScript, CSS Animation, Three.js, Design Systems, Static Deployment"),
        ("Tools", "Figma, Adobe Suite, Ableton / FL, Claude / Gemini, Git / GitHub"),
        ("Process", "Research, Prototyping, Client Direction, Production, Media Curation"),
    ]
    for row, (label, value) in zip(skills_table.rows, skills):
        for cell in row.cells:
            set_cell_border(cell, color="D8D4CC", size="4")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PAPER)
        row.cells[0].paragraphs[0].add_run(label.upper())
        set_run_font(row.cells[0].paragraphs[0].runs[0], size=7.6, color=RED, bold=True)
        row.cells[1].paragraphs[0].add_run(value)
        set_run_font(row.cells[1].paragraphs[0].runs[0], size=8.3, color=INK)

    add_section_title(doc, "Education / Training")
    add_role(doc, "Classically Trained Artist", "Charlotte, NC / Los Angeles, CA")
    add_bullet(doc, "[School / program / certification], [degree or focus], [year].")
    add_bullet(doc, "[Optional: awards, exhibitions, selected clients, press, or residencies.]")

    doc.save(RESUME_PATH)


def add_letter_body_paragraph(doc, text, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def build_cover_letter():
    doc = Document()
    section = set_document_base(doc, compact=False)
    add_footer(section)

    doc.core_properties.title = "Raehdojhn Cover Letter Template"
    doc.core_properties.subject = "Portfolio-compatible cover letter template"
    doc.core_properties.author = "Raehdojhn"

    add_nameplate(doc, "Cover Letter Template / Creative Technology / Design Systems")
    add_contact_strip(doc)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run("[Month Day, Year]\n[Hiring Manager Name]\n[Company / Studio]\n[Role Title]")
    set_run_font(run, size=9.2, color=MUTED)

    greeting = doc.add_paragraph()
    greeting.paragraph_format.space_after = Pt(8)
    run = greeting.add_run("Dear [Hiring Manager Name],")
    set_run_font(run, size=10.5, color=INK, bold=True)

    add_letter_body_paragraph(
        doc,
        "I am writing to express interest in the [Role Title] position at [Company]. "
        "My portfolio, Raehdojhn, is built as a working creative system: part brand archive, part interactive product, "
        "and part proof that I can move from concept to execution across design, code, audio, and AI-integrated workflows.",
    )
    add_letter_body_paragraph(
        doc,
        "I am a classically trained artist from Charlotte, North Carolina, now creating from Los Angeles. "
        "My work sits at the intersection of traditional craft and emerging technology. I build brand worlds, editorial interfaces, "
        "interactive project systems, music-player experiences, and deployable web environments with a strong sense of taste, structure, and cultural context.",
    )
    add_letter_body_paragraph(
        doc,
        "For [Company], I would bring concept-to-system thinking, production-ready design execution, and the ability to translate references into usable tools. "
        "Recent work includes the V4 Archive Site, a Mac-desktop-inspired portfolio environment with draggable windows, an iPod-style music player, Prompt Flux, "
        "and a media log; Future Archives - Ranaan x CCC, an interactive flip magazine; and Nodear*, an alter-ego brand and EPK portal combining identity, sound, and archive logic.",
    )

    add_section_title(doc, "Replace this section with role-specific proof")
    for bullet in [
        "If the role emphasizes brand: mention identity systems, visual direction, typography, campaign worlds, and art direction.",
        "If the role emphasizes product/design: mention UX, design systems, front-end implementation, prototyping, and interaction details.",
        "If the role emphasizes creative technology: mention JavaScript, CSS motion, AI workflows, media systems, and deployable web builds.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bullet)
        set_run_font(run, size=9.5, color=INK)

    add_letter_body_paragraph(
        doc,
        "I would be excited to bring that range to [Company] - not as disconnected disciplines, but as one coherent practice: creative direction that understands production, "
        "technical design that understands culture, and visual systems that can actually ship.",
    )
    add_letter_body_paragraph(
        doc,
        "Thank you for your time and consideration. I would welcome the chance to discuss how my portfolio and working process could support [Company / Team / Project].",
        after=12,
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run("Sincerely,\nRaehdojhn")
    set_run_font(run, size=10.5, color=INK, bold=True)

    doc.save(COVER_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_resume()
    build_cover_letter()
    print(RESUME_PATH)
    print(COVER_PATH)


if __name__ == "__main__":
    main()
